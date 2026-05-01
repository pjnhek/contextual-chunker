from pathlib import Path
from typing import Dict, Tuple

from docx import Document


def extract_docx(path: Path) -> Tuple[str, Dict]:
    """Extract text from a .docx file (paragraphs + tables)."""
    doc = Document(str(path))

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    # Don't strip the joined output — keeping trailing whitespace preserves
    # tokenization parity with downstream chunkers.
    return "\n".join(parts), {
        "source_file": str(path),
        "doc_type": "docx",
        "extractor": "python-docx",
    }
