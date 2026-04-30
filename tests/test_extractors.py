from pathlib import Path

import pytest

from contextual_chunker.extractors import extract_text


def test_txt_roundtrip(tmp_path: Path):
    p = tmp_path / "sample.txt"
    p.write_text("hello world\nsecond line\n", encoding="utf-8")

    text, meta = extract_text(p)

    assert "hello world" in text
    assert meta["doc_type"] == "txt"
    assert meta["extractor"] == "text"
    assert meta["source_file"] == str(p)


def test_md_roundtrip(tmp_path: Path):
    p = tmp_path / "notes.md"
    p.write_text("# Heading\n\nSome content here.", encoding="utf-8")

    text, meta = extract_text(p)

    assert "Heading" in text
    assert meta["doc_type"] == "md"


def test_pdf_extraction(tmp_path: Path):
    """Build a tiny PDF on the fly and round-trip the text out."""
    pytest.importorskip("fitz")
    import fitz

    p = tmp_path / "tiny.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello from a PDF")
    doc.save(p)
    doc.close()

    text, meta = extract_text(p)

    assert "Hello from a PDF" in text
    assert meta["doc_type"] == "pdf"
    assert meta["extractor"] == "pymupdf"
    assert meta["page_count"] == 1


def test_docx_extraction(tmp_path: Path):
    pytest.importorskip("docx")
    from docx import Document

    p = tmp_path / "tiny.docx"
    doc = Document()
    doc.add_paragraph("Para one.")
    doc.add_paragraph("Para two.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    doc.save(p)

    text, meta = extract_text(p)

    assert "Para one." in text
    assert "Para two." in text
    assert "A | B" in text
    assert meta["doc_type"] == "docx"


def test_unsupported_extension(tmp_path: Path):
    p = tmp_path / "weird.xyz"
    p.write_text("nope", encoding="utf-8")

    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text(p)
